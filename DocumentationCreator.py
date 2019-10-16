# Import Necessary packages
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

# Constant File Extensions to use as text
docx_Extension = ".docx"
png_Extenstion = ".png"

def MakeDocument(Title,plot,totalPoint,availablePoints):
    document = Document()
    section = document.sections[-1]
    newWidth, newHeight = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = newWidth
    section.page_height = newHeight
    document.add_heading(Title)
    document.add_picture(plot, width=Inches(10), height=Inches(5))
    document.add_page_break()
    # section.orientation = WD_ORIENT.PORTRAIT
    # section.page_width = newHeight
    # section.page_height = newWidth
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Total #Points"
    table.rows[0].cells[1].text = "Available #Points"
    table.rows[1].cells[0].text = str(totalPoint)
    table.rows[1].cells[1].text = str(availablePoints)
    document.save("Tryout.docx")


# Ideas to add to documentation creator:
# save display trendLine data
# have the CBDV logo as a Watermark/Logo
# Markus' signature. credibility
# WISHLIST:: All corresponding error plot and analysis (in a simple to read table??)
## Name the documents the title name : with some indication of date and time to never have conflicting file names
